from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(sys.argv[1])
OUT.parent.mkdir(parents=True, exist_ok=True)


def font(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(4)
run = title.add_run("SUPPLIER QUALIFICATION FORM")
font(run, 23, True, "0B2545")
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(16)
run = subtitle.add_run("Company information and declaration")
font(run, 13, False, "555555")

intro = doc.add_paragraph()
font(intro.add_run("Complete the fields below. Required fields are marked with *."), 10, False, "555555")

rows = [
    ("Legal company name *", "{{COMPANY_NAME}}"),
    ("Registration number *", "{{REGISTRATION_NUMBER}}"),
    ("VAT identification", "{{VAT_NUMBER}}"),
    ("Registered address *", "{{ADDRESS}}"),
    ("Primary contact *", "{{PRIMARY_CONTACT}}"),
    ("Contact email *", "{{CONTACT_EMAIL}}"),
    ("General liability policy *", "{{GL_POLICY_NUMBER}}"),
]
table = doc.add_table(rows=len(rows), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.autofit = False
table.style = "Table Grid"
for row, (label, placeholder) in zip(table.rows, rows):
    row.cells[0].width = Inches(1.875)
    row.cells[1].width = Inches(4.625)
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_margins(cell)
    p = row.cells[0].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    font(p.add_run(label), 10.5, True, "1F4D78")
    p = row.cells[1].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    font(p.add_run(placeholder), 10.5, False, "000000")

h = doc.add_paragraph()
h.paragraph_format.space_before = Pt(16)
h.paragraph_format.space_after = Pt(8)
font(h.add_run("Declaration"), 16, True, "2E74B5")
p = doc.add_paragraph()
font(p.add_run("I certify that the information supplied in this form is complete and accurate."), 11)

signature = doc.add_table(rows=2, cols=2)
signature.alignment = WD_TABLE_ALIGNMENT.LEFT
signature.autofit = False
sig_rows = [("Authorized signatory", "{{AUTHORIZED_SIGNATURE}}"), ("Date", "{{SIGNATURE_DATE}}")]
for row, (label, placeholder) in zip(signature.rows, sig_rows):
    row.cells[0].width = Inches(1.875)
    row.cells[1].width = Inches(4.625)
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_margins(cell)
    font(row.cells[0].paragraphs[0].add_run(label), 10.5, True, "1F4D78")
    font(row.cells[1].paragraphs[0].add_run(placeholder), 10.5)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(footer.add_run("Document filling end-to-end fixture"), 8.5, False, "777777")

doc.save(OUT)
